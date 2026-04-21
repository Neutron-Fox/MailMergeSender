"""
File import service for Universal Email Sender application.
Handles multi-format file import with error handling and threading support.
"""
import os
import re
from typing import Dict, Any, List, Tuple
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError

from .constants import (
    FILE_ENCODING_DEFAULT,
    FILE_ENCODING_FALLBACK,
    TEXT_FILE_DELIMITERS,
    TEXT_FILE_DELIMITER_FALLBACK,
    THREAD_MAX_WORKERS,
    IMPORT_THREAD_TIMEOUT
)
from .exceptions import (
    FileImportError,
    FileNotFoundError as CustomFileNotFoundError,
    UnsupportedFileTypeError,
    FileReadError,
    InvalidFileFormatError,
    NoDataFoundError,
    OperationTimeoutError
)
from .validators import FileValidator
from .logger_setup import get_logger

logger = get_logger(__name__)


class ImportService:
    """Handles file import operations for multiple file formats"""
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """
        Detect file type from extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            File type string
            
        Raises:
            UnsupportedFileTypeError: If file type not supported
        """
        try:
            ext = os.path.splitext(file_path.lower())[1]
            
            type_map = {
                '.xlsx': 'excel', '.xls': 'excel',
                '.docx': 'word', '.doc': 'word',
                '.csv': 'csv',
                '.txt': 'txt'
            }
            
            if ext not in type_map:
                raise UnsupportedFileTypeError(ext)
            
            file_type = type_map[ext]
            logger.info(f"Detected file type: {file_type} ({ext})")
            return file_type
        
        except UnsupportedFileTypeError:
            raise
        except Exception as e:
            logger.error(f"Error detecting file type: {e}")
            raise FileImportError(f"Error detecting file type: {e}")
    
    @staticmethod
    def import_file(file_path: str, timeout: int = None) -> Dict[str, Any]:
        """
        Import file using appropriate handler with timeout.
        
        Args:
            file_path: Path to file to import
            timeout: Maximum execution time in seconds
            
        Returns:
            Dictionary with import results
            
        Raises:
            FileImportError: If import fails
            OperationTimeoutError: If import exceeds timeout
        """
        try:
            logger.info(f"Starting file import: {file_path}")
            
            # Validate file
            FileValidator.validate_exists(file_path)
            FileValidator.validate_is_file(file_path)
            FileValidator.validate_readable(file_path)
            
            # Detect file type
            file_type = ImportService.detect_file_type(file_path)
            
            # Get appropriate handler
            handlers = {
                'excel': ImportService.read_excel_file,
                'word': ImportService.read_word_file,
                'csv': ImportService.read_csv_file,
                'txt': ImportService.read_txt_file
            }
            
            handler = handlers[file_type]
            
            # Run with timeout if specified
            if timeout:
                with ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(handler, file_path)
                    try:
                        result = future.result(timeout=timeout)
                    except FutureTimeoutError:
                        logger.error(f"Import timeout after {timeout}s")
                        raise OperationTimeoutError("File Import", timeout)
            else:
                result = handler(file_path)
            
            if result['success']:
                logger.info(
                    f"✓ File imported successfully: {result['message']} "
                    f"({len(result['data'])} rows)"
                )
            else:
                logger.error(f"✗ Import failed: {result['message']}")
            
            return result
        
        except (OperationTimeoutError, FileImportError):
            raise
        except Exception as e:
            logger.error(f"Unexpected error during import: {e}")
            raise FileImportError(f"Unexpected error: {e}")
    
    @staticmethod
    def read_excel_file(file_path: str) -> Dict[str, Any]:
        """
        Read Excel file (.xlsx, .xls).
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with headers, data, success, message
        """
        try:
            logger.debug(f"Reading Excel file: {file_path}")
            import pandas as pd
            
            try:
                df = pd.read_excel(file_path, engine='openpyxl')
            except Exception:
                logger.warning("openpyxl failed, trying xlrd...")
                df = pd.read_excel(file_path, engine='xlrd')
            
            if df.empty:
                raise NoDataFoundError(file_path)
            
            headers = list(df.columns)
            data = df.values.tolist()
            
            return {
                'headers': headers,
                'data': data,
                'success': True,
                'message': f'Excel file loaded: {len(df)} rows, {len(headers)} columns'
            }
        
        except NoDataFoundError:
            raise
        except ImportError:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': 'pandas/openpyxl not installed. Install with: pip install pandas openpyxl'
            }
        except Exception as e:
            logger.error(f"Error reading Excel file: {e}")
            raise FileReadError(file_path, str(e))
    
    @staticmethod
    def read_word_file(file_path: str) -> Dict[str, Any]:
        """
        Read Word file (.docx, .doc) - extracts tables or paragraphs.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary with headers, data, success, message
        """
        try:
            logger.debug(f"Reading Word file: {file_path}")
            import docx
            
            doc = docx.Document(file_path)
            tables_data = []
            
            # Try to extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []
                
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    
                    if i == 0:
                        headers = row_data
                    else:
                        table_data.append(row_data)
                
                if headers and table_data:
                    tables_data.append({
                        'headers': headers,
                        'data': table_data,
                        'table_number': table_idx
                    })
            
            # If tables found, use first table
            if tables_data:
                best_table = max(tables_data, key=lambda t: len(t['data']))
                return {
                    'headers': best_table['headers'],
                    'data': best_table['data'],
                    'success': True,
                    'message': f'Word file loaded: {len(best_table["data"])} rows from table'
                }
            
            # Fall back to paragraph text
            text_lines = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_lines.append([text])
            
            if text_lines:
                return {
                    'headers': ['Content'],
                    'data': text_lines,
                    'success': True,
                    'message': f'Word file loaded as text: {len(text_lines)} lines'
                }
            
            raise NoDataFoundError(file_path)
        
        except NoDataFoundError:
            raise
        except ImportError:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': 'python-docx not installed. Install with: pip install python-docx'
            }
        except Exception as e:
            logger.error(f"Error reading Word file: {e}")
            raise FileReadError(file_path, str(e))
    
    @staticmethod
    def read_csv_file(file_path: str) -> Dict[str, Any]:
        """
        Read CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Dictionary with headers, data, success, message
        """
        try:
            logger.debug(f"Reading CSV file: {file_path}")
            import pandas as pd
            
            df = pd.read_csv(file_path, encoding=FILE_ENCODING_DEFAULT)
            
            if df.empty:
                raise NoDataFoundError(file_path)
            
            headers = list(df.columns)
            data = df.values.tolist()
            
            return {
                'headers': headers,
                'data': data,
                'success': True,
                'message': f'CSV file loaded: {len(df)} rows, {len(headers)} columns'
            }
        
        except NoDataFoundError:
            raise
        except ImportError:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': 'pandas not installed. Install with: pip install pandas'
            }
        except Exception as e:
            logger.error(f"Error reading CSV file: {e}")
            raise FileReadError(file_path, str(e))
    
    @staticmethod
    def read_txt_file(file_path: str) -> Dict[str, Any]:
        """
        Read text file with smart delimiter detection.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with headers, data, success, message
        """
        try:
            logger.debug(f"Reading text file: {file_path}")
            
            # Try different encodings
            content = None
            for encoding in FILE_ENCODING_FALLBACK:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.debug(f"Successfully read file with encoding: {encoding}")
                    break
                except (UnicodeDecodeError, LookupError):
                    continue
            
            if content is None:
                raise InvalidFileFormatError(file_path, "Cannot decode file with any encoding")
            
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            
            if not lines:
                raise NoDataFoundError(file_path)
            
            # Detect delimiter
            delimiter = ImportService._detect_delimiter(lines)
            logger.debug(f"Detected delimiter: {repr(delimiter)}")
            
            # Parse with delimiter
            if delimiter and delimiter in TEXT_FILE_DELIMITERS:
                headers = [h.strip() for h in lines[0].split(delimiter)]
                data = []
                
                for line in lines[1:]:
                    row = [col.strip() for col in line.split(delimiter)]
                    # Pad missing columns
                    while len(row) < len(headers):
                        row.append('')
                    # Trim excess columns
                    row = row[:len(headers)]
                    data.append(row)
                
                return {
                    'headers': headers,
                    'data': data,
                    'success': True,
                    'message': f'Text file loaded: {len(data)} rows with delimiter "{delimiter}"'
                }
            else:
                # Single column
                return {
                    'headers': ['Content'],
                    'data': [[line] for line in lines],
                    'success': True,
                    'message': f'Text file loaded as single column: {len(lines)} lines'
                }
        
        except (NoDataFoundError, InvalidFileFormatError):
            raise
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise FileReadError(file_path, str(e))
    
    @staticmethod
    def _detect_delimiter(lines: List[str], delimiters: List[str] = None) -> str:
        """
        Detect delimiter in text file by analyzing first line.
        
        Args:
            lines: List of text lines
            delimiters: Delimiters to test (default: comma, tab, semicolon, pipe)
            
        Returns:
            Detected delimiter string
        """
        if not delimiters:
            delimiters = TEXT_FILE_DELIMITERS
        
        if not lines:
            return TEXT_FILE_DELIMITER_FALLBACK
        
        first_line = lines[0]
        best_delimiter = TEXT_FILE_DELIMITER_FALLBACK
        max_columns = 1
        
        for delimiter in delimiters:
            columns = len(first_line.split(delimiter))
            if columns > max_columns:
                max_columns = columns
                best_delimiter = delimiter
        
        return best_delimiter if max_columns > 1 else TEXT_FILE_DELIMITER_FALLBACK
