import os
import sys
import logging
import re
import subprocess
import time
from typing import List, Dict, Any
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QLabel, QPushButton, QFileDialog, QLineEdit, QMessageBox, QTextEdit, 
    QGroupBox, QTableWidget, QTableWidgetItem, QTabWidget, QComboBox, QProgressBar
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont
from theme import var_theme, get_button_style, get_table_style
logger = logging.getLogger(__name__)
class FileImporter:
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        ext = os.path.splitext(file_path.lower())[1]
        type_map = {
            '.xlsx': 'excel', '.xls': 'excel',
            '.docx': 'word', '.doc': 'word',
            '.csv': 'csv', '.txt': 'txt'
        }
        return type_map.get(ext, 'unknown')
    @staticmethod
    def read_excel_file(file_path: str) -> Dict[str, Any]:
        try:
            import pandas as pd
            df = pd.read_excel(file_path, engine='openpyxl')
            return {
                'headers': list(df.columns),
                'data': df.values.tolist(),
                'success': True,
                'message': f'Excel file loaded successfully. {len(df)} rows found.'
            }
        except Exception as e:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': f'Error reading Excel file: {str(e)}'
            }
    @staticmethod
    def read_word_file(file_path: str) -> Dict[str, Any]:
        try:
            import docx
            doc = docx.Document(file_path)
            tables_data = []
            for table in doc.tables:
                table_data = []
                headers = []
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if i == 0:
                        headers = row_data
                    else:
                        table_data.append(row_data)
                if headers and table_data:
                    tables_data.append({'headers': headers, 'data': table_data})
            if tables_data:
                return {
                    'headers': tables_data[0]['headers'],
                    'data': tables_data[0]['data'],
                    'success': True,
                    'message': f'Word file loaded successfully. {len(tables_data)} tables found.'
                }
            else:
                text_lines = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_lines.append([paragraph.text.strip()])
                return {
                    'headers': ['Content'],
                    'data': text_lines,
                    'success': True,
                    'message': f'Word file loaded as text. {len(text_lines)} lines found.'
                }
        except Exception as e:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': f'Error reading Word file: {str(e)}'
            }
    @staticmethod
    def read_csv_file(file_path: str) -> Dict[str, Any]:
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return {
                'headers': list(df.columns),
                'data': df.values.tolist(),
                'success': True,
                'message': f'CSV file loaded successfully. {len(df)} rows found.'
            }
        except Exception as e:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': f'Error reading CSV file: {str(e)}'
            }
    @staticmethod
    def read_txt_file(file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = [line.strip() for line in file.readlines() if line.strip()]
            if not lines:
                return {
                    'headers': ['Content'], 'data': [], 'success': True,
                    'message': 'Text file is empty.'
                }
            delimiters = [',', '\t', ';', '|']
            best_delimiter = ','
            max_columns = 1
            for delimiter in delimiters:
                columns = len(lines[0].split(delimiter))
                if columns > max_columns:
                    max_columns = columns
                    best_delimiter = delimiter
            if max_columns > 1:
                headers = [h.strip() for h in lines[0].split(best_delimiter)]
                data = []
                for line in lines[1:]:
                    row_data = [d.strip() for d in line.split(best_delimiter)]
                    while len(row_data) < len(headers):
                        row_data.append('')
                    data.append(row_data[:len(headers)])
                return {
                    'headers': headers, 'data': data, 'success': True,
                    'message': f'Text file loaded. {len(data)} rows with delimiter "{best_delimiter}".'
                }
            else:
                return {
                    'headers': ['Content'],
                    'data': [[line] for line in lines],
                    'success': True,
                    'message': f'Text file loaded as single column. {len(lines)} lines found.'
                }
        except Exception as e:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': f'Error reading text file: {str(e)}'
            }
    @staticmethod
    def import_file(file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {
                'headers': [], 'data': [], 'success': False,
                'message': 'File does not exist.'
            }
        file_type = FileImporter.detect_file_type(file_path)
        importers = {
            'excel': FileImporter.read_excel_file,
            'word': FileImporter.read_word_file,
            'csv': FileImporter.read_csv_file,
            'txt': FileImporter.read_txt_file
        }
        importer = importers.get(file_type)
        if importer:
            return importer(file_path)
        else:
            return {
                'headers': [], 'data': [], 'success': False,
                'message': f'Unsupported file type: {file_type}'
            }
class PlaceholderExtractor:
    @staticmethod
    def extract_placeholders(text: str) -> List[str]:
        placeholders = set()
        patterns = [
            r'\{([^}]+)\}',        
            r'<([^>]+)>',          
            r'\[\[([^\]]+)\]\]',   
            r'<<([^>]+)>>',        
            r'\(([^)]+)\)',        
            r'\{\{([^}]+)\}\}',    
            r'\[([^\]]+)\]',       
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                match = match.strip().upper()
                if match and any(c.isalpha() for c in match):
                    standardized = '{' + match + '}'
                    placeholders.add(standardized)
        return sorted(list(placeholders))
    @staticmethod
    def suggest_mappings(placeholders: List[str], headers: List[str]) -> Dict[str, str]:
        suggestions = {}
        for placeholder in placeholders:
            placeholder = placeholder.strip('{}').upper()
            best_match = None
            best_score = 0
            for header in headers:
                header = header.upper()
                if placeholder == header:
                    best_match = header
                    best_score = 100
                    break
                if placeholder in header or header in placeholder:
                    score = 80
                    if score > best_score:
                        best_match = header
                        best_score = score
                common_mappings = {
                    ('NAME', 'FULLNAME', 'FULL_NAME'): ('NAME', 'PERSON', 'USER'),
                    ('EMAIL', 'MAIL', 'E_MAIL'): ('EMAIL', 'MAIL', '@')
                }
                for placeholder_group, header_keywords in common_mappings.items():
                    if placeholder in placeholder_group:
                        if any(keyword in header for keyword in header_keywords):
                            if 60 > best_score:
                                best_match = header
                                best_score = 60
            if best_match:
                suggestions[placeholder] = best_match
        return suggestions
class EmailSender:
    _outlook_instance = None
    @staticmethod
    def is_outlook_running() -> bool:
        """Check if Outlook process is running"""
        try:
            result = subprocess.run(['tasklist', '/FI', 'IMAGENAME eq OUTLOOK.EXE'], 
                                  capture_output=True, text=True, timeout=5)
            if 'OUTLOOK.EXE' in result.stdout:
                logger.info("Found Outlook process running")
                return True
        except Exception as e:
            logger.warning(f"Could not check if Outlook is running: {e}")
        logger.info("Outlook process not detected")
        return False
    @staticmethod
    def start_outlook():
        """Try to start Outlook application in background (minimized, no window popup)"""
        try:
            logger.info("Starting Microsoft Outlook in background...")
            outlook_paths = [
                r"C:\Program Files\Microsoft Office\root\Office16\OUTLOOK.EXE",
                r"C:\Program Files (x86)\Microsoft Office\root\Office16\OUTLOOK.EXE",
                r"C:\Program Files\Microsoft Office\Office16\OUTLOOK.EXE",
                r"C:\Program Files (x86)\Microsoft Office\Office16\OUTLOOK.EXE",
                r"C:\Program Files\Microsoft Office\root\Office15\OUTLOOK.EXE",
                r"C:\Program Files (x86)\Microsoft Office\root\Office15\OUTLOOK.EXE",
            ]
            for path in outlook_paths:
                if os.path.exists(path):
                    logger.info(f"Found Outlook at: {path}")
                    import subprocess
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                    startupinfo.wShowWindow = 6  
                    subprocess.Popen([path], startupinfo=startupinfo)
                    logger.info("Waiting for Outlook to start in background...")
                    wait_time = 8 if hasattr(sys, 'frozen') else 5
                    time.sleep(wait_time)
                    if EmailSender.is_outlook_running():
                        logger.info("✓ Outlook started successfully in background")
                        return True
                    else:
                        logger.warning("Outlook process not detected after start attempt")
                        time.sleep(2)  
                        return EmailSender.is_outlook_running()
            try:
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = 6  
                subprocess.Popen(["outlook.exe"], startupinfo=startupinfo)
                logger.info("Started Outlook via system PATH (minimized)")
                wait_time = 8 if hasattr(sys, 'frozen') else 5
                time.sleep(wait_time)
                return EmailSender.is_outlook_running()
            except Exception as e:
                logger.warning(f"Could not start Outlook via PATH: {e}")
            logger.warning("Could not find or start Outlook")
            return False
        except Exception as e:
            logger.error(f"Error starting Outlook: {e}")
            return False
    @staticmethod
    def get_email_accounts() -> List[Dict[str, Any]]:
        """Extract email accounts from Microsoft Outlook using pywin32"""
        accounts = []
        try:
            logger.info("Loading Outlook email accounts via pywin32...")
            logger.info("Connecting to Outlook silently (COM will start it in background if needed)...")
            try:
                import win32com.client
                logger.info("win32com.client imported successfully")
            except ImportError as e:
                logger.error(f"CRITICAL: Failed to import win32com.client: {e}")
                logger.error("Install pywin32: pip install pywin32")
                QMessageBox.critical(None, "Missing Dependency", 
                    "pywin32 package is required!\n\nInstall it with: pip install pywin32")
                return accounts
            logger.info("Connecting to Outlook Application...")
            outlook = None
            if EmailSender._outlook_instance is not None:
                try:
                    _ = EmailSender._outlook_instance.Version
                    logger.info("✓ Reusing existing Outlook instance")
                    outlook = EmailSender._outlook_instance
                except:
                    logger.warning("Cached Outlook instance is invalid, creating new connection")
                    EmailSender._outlook_instance = None
            if EmailSender._outlook_instance is None:
                try:
                    outlook = win32com.client.Dispatch("Outlook.Application")
                    _ = outlook.Version
                    logger.info(f"Successfully connected to Outlook (version {outlook.Version})")
                    EmailSender._outlook_instance = outlook
                    logger.info("✓ Outlook instance cached for reuse")
                except Exception as e:
                    logger.error(f"Failed to connect to Outlook: {e}")
                    error_msg = (
                        f"Could not connect to Microsoft Outlook.\n\n"
                        f"Error: {str(e)}\n\n"
                        "Solutions:\n"
                        "1. Check if Outlook has a security prompt asking for permission - allow it\n"
                        "2. Wait a few seconds for Outlook to fully start, then try again\n"
                        "3. Run this program as Administrator (right-click → Run as administrator)\n"
                        "4. Close ALL Outlook windows and restart this program"
                    )
                    QMessageBox.critical(None, "Outlook Connection Error", error_msg)
                    return accounts
            QApplication.processEvents()
            try:
                namespace = outlook.GetNamespace("MAPI")
                logger.info("Connected to MAPI namespace")
            except Exception as e:
                logger.error(f"Failed to get MAPI namespace: {e}")
                return accounts
            try:
                outlook_accounts = outlook.Session.Accounts
                account_count = outlook_accounts.Count
                logger.info(f"Found {account_count} Outlook account(s)")
                if account_count == 0:
                    logger.warning("No Outlook accounts configured!")
                    QMessageBox.warning(None, "No Accounts", 
                        "No email accounts found in Outlook.\n\n"
                        "Please configure at least one email account in Outlook.")
                    return accounts
            except Exception as e:
                logger.error(f"Failed to access Outlook accounts: {e}")
                return accounts
            for i in range(1, account_count + 1):
                try:
                    account = outlook_accounts.Item(i)
                    account_name = account.DisplayName
                    try:
                        email_address = account.SmtpAddress
                        if email_address and '@' in email_address:
                            list_index = len(accounts)  
                            accounts.append({
                                'email': email_address,
                                'account_object': account
                            })
                            logger.info(f"✓ Outlook Position {i} → List Index[{list_index}]: {email_address} (Name: {account_name})")
                        else:
                            logger.warning(f"✗ Account {i}: No valid SMTP address")
                    except Exception as e:
                        logger.warning(f"✗ Account {i}: Error: {e}")
                    QApplication.processEvents()
                except Exception as e:
                    logger.warning(f"Error processing account {i}: {e}")
        except Exception as e:
            logger.error(f"Critical error loading email accounts: {e}")
            import traceback
            logger.error(traceback.format_exc())
            QMessageBox.critical(None, "Error", 
                f"Error loading email accounts:\n\n{str(e)}\n\n"
                "Please ensure:\n"
                "1. Microsoft Outlook is installed\n"
                "2. pywin32 is installed (pip install pywin32)\n"
                "3. Outlook has configured email accounts")
        if len(accounts) == 0:
            logger.warning("No email accounts loaded!")
            print(f"\n⚠ WARNING: No email accounts found in Outlook!\n")
        else:
            logger.info(f"Successfully loaded {len(accounts)} email account(s)")
            print(f"\n✓ Successfully loaded {len(accounts)} email account(s):")
            for i, acc in enumerate(accounts, 1):
                print(f"  {i}. {acc['email']}")
            print()
        return accounts
    @staticmethod
    def send_emails(recipients: List[Dict], subject: str, template: str, 
                   account: Dict, attachments: List[str] = None) -> Dict[str, Any]:
        """Send emails using Microsoft Outlook via pywin32"""
        try:
            import win32com.client
            sender_email = account.get('email', None)
            if not sender_email:
                return {
                    'success': False,
                    'message': 'No valid account selected.',
                    'sent': 0,
                    'failed': len(recipients)
                }
            if EmailSender._outlook_instance is not None:
                try:
                    _ = EmailSender._outlook_instance.Version
                    outlook = EmailSender._outlook_instance
                    logger.info("✓ Reusing cached Outlook instance for sending")
                except:
                    logger.warning("Cached Outlook instance is invalid, will reconnect")
                    EmailSender._outlook_instance = None
            if EmailSender._outlook_instance is None:
                logger.info("No cached instance - connecting to Outlook")
                if not EmailSender.is_outlook_running():
                    logger.error("Outlook is not running - this should not happen at send time")
                    return {
                        'success': False,
                        'message': 'Outlook is not running. Please restart the application.',
                        'sent': 0,
                        'failed': len(recipients)
                    }
                try:
                    logger.info("Connecting to existing Outlook instance...")
                    outlook = win32com.client.Dispatch("Outlook.Application")
                    _ = outlook.Version
                    EmailSender._outlook_instance = outlook
                    logger.info("✓ Connected and cached Outlook instance")
                except Exception as e:
                    logger.error(f"Failed to connect to Outlook: {e}")
                    return {
                        'success': False,
                        'message': f'Cannot connect to Outlook: {str(e)}',
                        'sent': 0,
                        'failed': len(recipients)
                    }
            account_object = None
            try:
                outlook_accounts = outlook.Session.Accounts
                logger.info(f"Searching for account: {sender_email}")
                for i in range(1, outlook_accounts.Count + 1):
                    acc = outlook_accounts.Item(i)
                    acc_email = acc.SmtpAddress
                    logger.info(f"  Checking account {i}: {acc_email}")
                    if acc_email.lower() == sender_email.lower():
                        account_object = acc
                        logger.info(f"✓ FOUND MATCHING ACCOUNT: {acc_email}")
                        logger.info(f"  Account DisplayName: {acc.DisplayName}")
                        break
                if not account_object:
                    logger.error(f"✗ ACCOUNT NOT FOUND: {sender_email}")
                    return {
                        'success': False,
                        'message': f'Could not find account {sender_email} in Outlook session',
                        'sent': 0,
                        'failed': len(recipients)
                    }
            except Exception as e:
                logger.error(f"Error finding account: {e}")
                return {
                    'success': False,
                    'message': f'Error accessing Outlook accounts: {str(e)}',
                    'sent': 0,
                    'failed': len(recipients)
                }
            sent_count = 0
            failed_count = 0
            failed_recipients = []
            for i, recipient_data in enumerate(recipients, 1):
                try:
                    recipient_email = None
                    for field in ['EMAIL', 'Email', 'email', 'E-mail', 'E-Mail', 'Mail', 'MAIL']:
                        if field in recipient_data and recipient_data[field]:
                            recipient_email = str(recipient_data[field]).strip()
                            break
                    if not recipient_email or '@' not in recipient_email:
                        failed_count += 1
                        failed_recipients.append(f"Recipient {i}: No valid email")
                        continue
                    mail_item = outlook.CreateItem(0)  
                    mail_item.SendUsingAccount = account_object
                    try:
                        mail_item.SentOnBehalfOfName = sender_email
                        logger.info(f"Email {i}: SentOnBehalfOfName set to: {sender_email}")
                    except Exception as e:
                        logger.warning(f"Email {i}: Could not set SentOnBehalfOfName: {e}")
                    logger.info(f"Email {i}: Account set immediately after creation: {sender_email}")
                    mail_item.To = recipient_email
                    if '_processed_subject' in recipient_data:
                        mail_item.Subject = recipient_data['_processed_subject']
                    else:
                        mail_item.Subject = subject
                    if '_processed_template' in recipient_data:
                        body_text = recipient_data['_processed_template']
                    else:
                        body_text = template
                    mail_item.HTMLBody = body_text.replace('\n', '<br>')
                    if attachments:
                        for att_path in attachments:
                            if os.path.exists(att_path):
                                try:
                                    mail_item.Attachments.Add(att_path)
                                except:
                                    pass
                    logger.info(f"Email {i}: Setting sender account to: {sender_email}")
                    mail_item.SendUsingAccount = account_object
                    mail_item.Save()
                    logger.info(f"Email {i}: Email saved")
                    mail_item.SendUsingAccount = account_object
                    time.sleep(0.05)
                    mail_item.SendUsingAccount = account_object
                    try:
                        test_sender = mail_item.SendUsingAccount
                        if test_sender:
                            logger.info(f"Email {i}: Final sender check: {test_sender.SmtpAddress}")
                        else:
                            logger.warning(f"Email {i}: SendUsingAccount returned None (Outlook quirk)")
                            mail_item.SendUsingAccount = account_object
                    except Exception as e:
                        logger.warning(f"Email {i}: Could not verify sender: {e}")
                    mail_item.Send()
                    logger.info(f"Email {i}: ✓ Sent from: {sender_email}")
                    time.sleep(0.1)
                    sent_count += 1
                except Exception as e:
                    failed_count += 1
                    error_msg = f"Recipient {i}: {str(e)}"
                    failed_recipients.append(error_msg)
            return {
                'success': failed_count == 0,
                'message': f'Sent {sent_count} emails' + (f', {failed_count} failed' if failed_count > 0 else ''),
                'sent': sent_count,
                'failed': failed_count,
                'failed_details': failed_recipients
            }
        except Exception as e:
            return {
                'success': False,
                'message': f'Critical error: {str(e)}',
                'sent': 0,
                'failed': len(recipients)
            }
    @staticmethod
    def _replace_placeholders(text: str, data: Dict[str, Any]) -> str:
        result = text
        for key, value in data.items():
            placeholders = [
                f'{{{key.upper()}}}',
                f'<{key.upper()}>',
                f'[{key.upper()}]',
                f'{{{{{key.upper()}}}}}',
                f'<<{key.upper()}>>',
                f'[[{key.upper()}]]'
            ]
            str_value = str(value) if value is not None else ''
            for placeholder in placeholders:
                result = result.replace(placeholder, str_value)
        return result
class UniversalSender(QMainWindow):
    def __init__(self, loading_screen=None):
        super().__init__()
        self.loading_screen = loading_screen
        self.setWindowTitle("Universal Email Sender")
        self.setGeometry(100, 100, 1200, 800)  
        self.setMinimumSize(1200, 800)
        self.setMaximumSize(1200, 800)  
        self.imported_data = []
        self.processed_data = []
        self.filtered_data = []
        self.selected_rows = set()  
        self.attachments = []  
        self.email_accounts_list = []  
        self.headers = []
        self.placeholders = []
        self.column_mapping = {}
        self.email_accounts = []
        self.template_formatting = {}  
        self.bullet_styles = {
            "Dash": "-",
            "Bullet": "•",
            "Circle": "o",
            "Arrow": "→",
            "Star": "★"
        }
        self.tab_widgets = {}
        self.tabs_created = set()
        self.email_accounts_loaded = False
        self.replacement_pairs = []  
        self.setup_ui()
        self.apply_theme()
    def setup_ui(self):
        """Setup the main UI and pre-load all tabs"""
        self.apply_dark_titlebar()
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(8)
        header_label = QLabel("Universal Email Sender")
        header_label.setAlignment(Qt.AlignCenter)
        header_label.setFont(var_theme.get_font(18, 'bold'))
        header_label.setFixedHeight(35)
        header_label.setStyleSheet(f"color: {var_theme.colors['button_primary']};")
        layout.addWidget(header_label)
        self.tabs = QTabWidget()
        self.tabs.setStyleSheet(self.get_clickable_tab_style())
        self.tabs.currentChanged.connect(self.on_tab_changed)
        self.load_all_tabs()
        layout.addWidget(self.tabs)
        self.statusBar().showMessage("Ready - Import a file to get started")
    def create_import_tab(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(12)
        layout.setContentsMargins(15, 15, 15, 15)
        file_group = QGroupBox("Select File to Import")
        file_layout = QVBoxLayout()
        file_layout.setContentsMargins(12, 12, 12, 12)
        file_layout.setSpacing(10)
        path_layout = QHBoxLayout()
        self.file_path_input = QLineEdit()
        self.file_path_input.setPlaceholderText("Select file to import...")
        browse_btn = QPushButton("Browse")
        browse_btn.setStyleSheet(get_button_style('primary'))
        browse_btn.clicked.connect(self.browse_file)
        path_layout.addWidget(QLabel("File:"))
        path_layout.addWidget(self.file_path_input)
        path_layout.addWidget(browse_btn)
        file_layout.addLayout(path_layout)
        import_layout = QHBoxLayout()
        import_layout.addStretch()
        self.import_btn = QPushButton("Import File")
        self.import_btn.setStyleSheet(get_button_style('success'))
        self.import_btn.clicked.connect(self.import_file)
        self.import_btn.setEnabled(False)
        import_layout.addWidget(self.import_btn)
        import_layout.addStretch()
        file_layout.addLayout(import_layout)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        preview_group = QGroupBox("Imported Data Preview")
        preview_layout = QVBoxLayout()
        preview_layout.setContentsMargins(12, 15, 12, 12)
        preview_layout.setSpacing(10)
        search_filter_layout = QHBoxLayout()
        select_all_btn = QPushButton("Select All")
        select_all_btn.setStyleSheet(get_button_style('primary'))
        select_all_btn.setMaximumWidth(100)
        select_all_btn.clicked.connect(self.select_all_rows)
        deselect_all_btn = QPushButton("Deselect All")
        deselect_all_btn.setStyleSheet(get_button_style('default'))
        deselect_all_btn.setMaximumWidth(100)
        deselect_all_btn.clicked.connect(self.deselect_all_rows)
        search_filter_layout.addWidget(select_all_btn)
        search_filter_layout.addWidget(deselect_all_btn)
        search_filter_layout.addWidget(QLabel("|"))  
        search_label = QLabel("Search:")
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Type to search in all columns...")
        self.search_input.setMaximumWidth(250)
        self.search_input.textChanged.connect(self.filter_table_data)
        filter_label = QLabel("Sort Column:")
        self.filter_column_combo = QComboBox()
        self.filter_column_combo.setMaximumWidth(150)
        self.filter_column_combo.currentTextChanged.connect(self.sort_table_data)
        sort_order_label = QLabel("Order:")
        self.sort_order_combo = QComboBox()
        self.sort_order_combo.addItems(["A-Z", "Z-A"])
        self.sort_order_combo.setMaximumWidth(80)
        self.sort_order_combo.currentTextChanged.connect(self.sort_table_data)
        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet(get_button_style('default'))
        clear_btn.setMaximumWidth(80)
        clear_btn.clicked.connect(self.clear_filters)
        search_filter_layout.addWidget(search_label)
        search_filter_layout.addWidget(self.search_input)
        search_filter_layout.addWidget(filter_label)
        search_filter_layout.addWidget(self.filter_column_combo)
        search_filter_layout.addWidget(sort_order_label)
        search_filter_layout.addWidget(self.sort_order_combo)
        search_filter_layout.addWidget(clear_btn)
        search_filter_layout.addStretch()
        preview_layout.addLayout(search_filter_layout)
        self.data_table = QTableWidget()
        self.data_table.setMinimumHeight(200)
        self.data_table.setMaximumHeight(300)
        self.data_table.setStyleSheet(get_table_style())
        self.data_table.setAlternatingRowColors(True)
        self.data_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.data_table.setSelectionMode(QTableWidget.MultiSelection)
        preview_layout.addWidget(self.data_table)
        self.selection_info_label = QLabel("No data loaded")
        self.selection_info_label.setStyleSheet(f"color: {var_theme.colors['text_muted']}; font-size: 9pt; padding: 5px;")
        preview_layout.addWidget(self.selection_info_label)
        nav_layout = QHBoxLayout()
        nav_layout.addStretch()
        self.next_btn_1 = QPushButton("Next: Compose Email →")
        self.next_btn_1.setStyleSheet(get_button_style('primary'))
        self.next_btn_1.setMinimumHeight(32)
        self.next_btn_1.setMinimumWidth(140)
        self.next_btn_1.clicked.connect(lambda: self.tabs.setCurrentIndex(1))
        self.next_btn_1.setEnabled(False)
        nav_layout.addWidget(self.next_btn_1)
        preview_layout.addLayout(nav_layout)
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group, 1)
        return widget
    def create_compose_tab(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(10)
        layout.setContentsMargins(15, 15, 15, 15)
        top_section_layout = QHBoxLayout()
        subject_group = QGroupBox("Email Subject")
        subject_layout = QVBoxLayout()
        subject_layout.setContentsMargins(8, 8, 8, 8)
        subject_layout.setSpacing(4)
        self.subject_input = QLineEdit()
        self.subject_input.setPlaceholderText("Enter email subject (use {NAME}, {EMAIL}, {DEVICES}, etc.)")
        self.subject_input.textChanged.connect(self.update_send_summary)
        self.subject_input.setMinimumHeight(28)
        subject_layout.addWidget(self.subject_input)
        subject_group.setLayout(subject_layout)
        top_section_layout.addWidget(subject_group, 1)  
        attachments_group = QGroupBox("Attachments")
        attachments_layout = QVBoxLayout()
        attachments_layout.setContentsMargins(8, 8, 8, 8)
        attachments_layout.setSpacing(4)
        attachments_buttons_layout = QHBoxLayout()
        attachments_buttons_layout.setSpacing(5)  
        add_attachment_btn = QPushButton("Add File")
        add_attachment_btn.setStyleSheet(get_button_style('primary'))
        add_attachment_btn.setMinimumSize(70, 26)
        add_attachment_btn.clicked.connect(self.add_attachment)
        remove_attachment_btn = QPushButton("Remove")
        remove_attachment_btn.setStyleSheet(get_button_style('danger'))
        remove_attachment_btn.setMinimumSize(70, 26)
        remove_attachment_btn.clicked.connect(self.remove_attachment)
        clear_attachments_btn = QPushButton("Clear All")
        clear_attachments_btn.setStyleSheet(get_button_style('default'))
        clear_attachments_btn.setMinimumSize(70, 26)
        clear_attachments_btn.clicked.connect(self.clear_attachments)
        attachments_buttons_layout.addWidget(add_attachment_btn)
        attachments_buttons_layout.addWidget(remove_attachment_btn)
        attachments_buttons_layout.addWidget(clear_attachments_btn)
        attachments_buttons_layout.addStretch()
        attachments_layout.addLayout(attachments_buttons_layout)
        self.attachments_table = QTableWidget()
        self.attachments_table.setColumnCount(2)
        self.attachments_table.setHorizontalHeaderLabels(["File", "Size"])
        self.attachments_table.setMinimumHeight(100)
        self.attachments_table.setMaximumHeight(130)
        self.attachments_table.setStyleSheet(get_table_style())
        self.attachments_table.setAlternatingRowColors(True)
        self.attachments_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.attachments_table.horizontalHeader().setStretchLastSection(True)
        self.attachments_table.setColumnWidth(0, 250)
        self.attachments_table.verticalHeader().setVisible(False)
        attachments_layout.addWidget(self.attachments_table)
        self.attachment_info_label = QLabel("No files")
        self.attachment_info_label.setStyleSheet(f"color: {var_theme.colors['text_muted']}; font-size: 8pt; padding: 2px;")
        attachments_layout.addWidget(self.attachment_info_label)
        attachments_group.setLayout(attachments_layout)
        top_section_layout.addWidget(attachments_group, 1)  
        layout.addLayout(top_section_layout)
        template_group = QGroupBox("Email Template")
        template_layout = QVBoxLayout()
        template_layout.setContentsMargins(12, 12, 12, 12)
        template_layout.setSpacing(8)
        template_btn_layout = QHBoxLayout()
        load_template_btn = QPushButton("Load")
        load_template_btn.setStyleSheet(get_button_style('default'))
        load_template_btn.setMaximumWidth(60)
        load_template_btn.setMaximumHeight(30)
        load_template_btn.clicked.connect(self.load_template)
        save_template_btn = QPushButton("Save")
        save_template_btn.setStyleSheet(get_button_style('default'))
        save_template_btn.setMaximumWidth(60)
        save_template_btn.setMaximumHeight(30)
        save_template_btn.clicked.connect(self.save_template)
        detect_btn = QPushButton("Detect Placeholders")
        detect_btn.setStyleSheet(get_button_style('success'))
        detect_btn.setMaximumWidth(130)
        detect_btn.setMaximumHeight(30)
        detect_btn.clicked.connect(self.detect_placeholders)
        template_btn_layout.addWidget(load_template_btn)
        template_btn_layout.addWidget(save_template_btn)
        template_btn_layout.addWidget(detect_btn)
        template_btn_layout.addStretch()
        template_layout.addLayout(template_btn_layout)
        self.template_editor = QTextEdit()
        self.template_editor.setMinimumHeight(200)
        self.template_editor.setMaximumHeight(250)
        self.template_editor.setFont(var_theme.get_font(10))
        self.template_editor.textChanged.connect(self.update_send_summary)
        self.template_editor.textChanged.connect(self.detect_placeholders)
        self.template_editor.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.template_editor.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.template_editor.setLineWrapMode(QTextEdit.WidgetWidth)
        template_layout.addWidget(self.template_editor)
        # No default template - user must load from file using the Load button
        self.placeholders_label = QLabel("Detected placeholders will appear here...")
        self.placeholders_label.setWordWrap(True)
        self.placeholders_label.setFixedHeight(50)  
        self.placeholders_label.setStyleSheet(f"color: {var_theme.colors['text_muted']}; padding: 8px; background-color: {var_theme.colors['secondary_bg']}; border-radius: 4px; font-size: 9pt;")
        template_layout.addWidget(self.placeholders_label)
        template_group.setLayout(template_layout)
        layout.addWidget(template_group, 1)  
        nav_layout = QHBoxLayout()
        back_btn_2 = QPushButton("← Back")
        back_btn_2.setStyleSheet(get_button_style('default'))
        back_btn_2.clicked.connect(lambda: self.tabs.setCurrentIndex(0))
        self.next_btn_2 = QPushButton("Next: Map Columns →")
        self.next_btn_2.setStyleSheet(get_button_style('primary'))
        self.next_btn_2.setMinimumHeight(32)
        self.next_btn_2.setMinimumWidth(150)
        self.next_btn_2.clicked.connect(lambda: self.tabs.setCurrentIndex(2))
        self.next_btn_2.setEnabled(False)
        nav_layout.addWidget(back_btn_2)
        nav_layout.addStretch()
        nav_layout.addWidget(self.next_btn_2)
        layout.addLayout(nav_layout)
        return widget
    def create_mapping_tab(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(12)
        layout.setContentsMargins(15, 15, 15, 15)
        mapping_group = QGroupBox("Column Mapping")
        mapping_layout = QVBoxLayout()
        mapping_layout.setContentsMargins(12, 12, 12, 12)
        mapping_layout.setSpacing(10)
        mapping_label = QLabel("Map data columns to email placeholders:")
        mapping_label.setFont(var_theme.get_font(11, 'medium'))
        mapping_layout.addWidget(mapping_label)
        instructions_label = QLabel("Select the data column that corresponds to each placeholder in your email template.")
        instructions_label.setStyleSheet(f"color: {var_theme.colors['text_muted']}; font-size: 9pt; margin-bottom: 10px;")
        instructions_label.setWordWrap(True)
        mapping_layout.addWidget(instructions_label)
        self.mapping_table = QTableWidget()
        self.mapping_table.setColumnCount(3)
        self.mapping_table.setHorizontalHeaderLabels(["Placeholder", "Data Column", "Sample Data"])
        self.mapping_table.setMinimumHeight(350)
        self.mapping_table.setMaximumHeight(400)  
        self.mapping_table.setStyleSheet(get_table_style())
        self.mapping_table.setAlternatingRowColors(True)
        self.mapping_table.horizontalHeader().setStretchLastSection(False)
        self.mapping_table.setColumnWidth(0, 150)  
        self.mapping_table.setColumnWidth(1, 200)  
        self.mapping_table.setColumnWidth(2, 250)  
        self.mapping_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        mapping_layout.addWidget(self.mapping_table)
        mapping_group.setLayout(mapping_layout)
        layout.addWidget(mapping_group)
        nav_layout = QHBoxLayout()
        back_btn_3 = QPushButton("← Back")
        back_btn_3.setStyleSheet(get_button_style('default'))
        back_btn_3.setMinimumHeight(32)
        back_btn_3.setMinimumWidth(80)
        back_btn_3.clicked.connect(lambda: self.tabs.setCurrentIndex(1))
        self.next_btn_3 = QPushButton("Next: Template Formatting →")
        self.next_btn_3.setStyleSheet(get_button_style('primary'))
        self.next_btn_3.setMinimumHeight(32)
        self.next_btn_3.setMinimumWidth(150)
        self.next_btn_3.clicked.connect(lambda: self.tabs.setCurrentIndex(3))
        nav_layout.addWidget(back_btn_3)
        nav_layout.addStretch()
        nav_layout.addWidget(self.next_btn_3)
        layout.addLayout(nav_layout)
        return widget
    def create_template_formatting_tab(self) -> QWidget:
        from PyQt5.QtWidgets import QScrollArea, QCheckBox
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setSpacing(6)
        main_layout.setContentsMargins(8, 8, 8, 8)
        column_group = QGroupBox("Step 1: Select Column to Format")
        column_layout = QVBoxLayout()
        column_layout.setContentsMargins(6, 6, 6, 6)
        column_layout.setSpacing(4)
        column_select_layout = QHBoxLayout()
        column_label = QLabel("Column:")
        column_label.setMinimumWidth(70)
        column_select_layout.addWidget(column_label)
        self.format_column_combo = QComboBox()
        self.format_column_combo.setMinimumWidth(200)
        self.format_column_combo.setMinimumHeight(24)
        self.format_column_combo.currentTextChanged.connect(self.load_formatting_settings)
        column_select_layout.addWidget(self.format_column_combo)
        column_select_layout.addStretch()
        column_layout.addLayout(column_select_layout)
        column_group.setLayout(column_layout)
        main_layout.addWidget(column_group)
        bullet_group = QGroupBox("Step 2: Bullet Point Formatting")
        bullet_layout = QVBoxLayout()
        bullet_layout.setContentsMargins(6, 6, 6, 6)
        bullet_layout.setSpacing(4)
        bullet_options_layout = QHBoxLayout()
        bullet_style_label = QLabel("Bullet style:")
        bullet_style_label.setMinimumWidth(70)
        bullet_options_layout.addWidget(bullet_style_label)
        self.enable_bullet_checkbox = QCheckBox("Enable bullet point formatting")
        self.enable_bullet_checkbox.stateChanged.connect(self.on_bullet_checkbox_changed)
        bullet_options_layout.addWidget(self.enable_bullet_checkbox)
        self.bullet_combo = QComboBox()
        for name, symbol in self.bullet_styles.items():
            self.bullet_combo.addItem(f"{name} ({symbol})", symbol)
        self.bullet_combo.setMinimumWidth(120)
        self.bullet_combo.setMaximumWidth(150)
        self.bullet_combo.setMinimumHeight(24)
        self.bullet_combo.setEnabled(False)
        self.bullet_combo.currentTextChanged.connect(self.auto_save_and_update_preview)
        bullet_options_layout.addWidget(self.bullet_combo)
        bullet_options_layout.addStretch()
        bullet_layout.addLayout(bullet_options_layout)
        bullet_group.setLayout(bullet_layout)
        main_layout.addWidget(bullet_group)
        replacement_section = QGroupBox("Step 3: String Replacements (Optional)")
        replacement_section_layout = QVBoxLayout()
        replacement_section_layout.setContentsMargins(6, 6, 6, 6)
        replacement_section_layout.setSpacing(6)
        two_box_layout = QHBoxLayout()
        two_box_layout.setSpacing(10)
        add_box = QGroupBox("Add New Replacement")
        add_box.setStyleSheet(f"""
            QGroupBox {{
                background-color: {var_theme.colors['secondary_bg']};
                border: 2px solid {var_theme.colors['border_primary']};
                border-radius: 8px;
                font-weight: 600;
                font-size: 10pt;
                color: {var_theme.colors['text_primary']};
                padding-top: 18px;
                margin-top: 12px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top center;
                background-color: {var_theme.colors['secondary_bg']};
                padding: 4px 12px;
                border: 1px solid {var_theme.colors['border_primary']};
                border-radius: 4px;
            }}
        """)
        add_box_layout = QVBoxLayout()
        add_box_layout.setContentsMargins(8, 10, 8, 8)
        add_box_layout.setSpacing(6)
        find_layout = QHBoxLayout()
        find_label = QLabel("Find:")
        find_label.setMinimumWidth(55)
        find_layout.addWidget(find_label)
        self.new_find_input = QLineEdit()
        self.new_find_input.setPlaceholderText("Text to find...")
        self.new_find_input.setMinimumHeight(24)
        self.new_find_input.setStyleSheet(f"""
            QLineEdit {{
                background-color: {var_theme.colors['input_bg']};
                border: 2px solid {var_theme.colors['border_primary']};
                border-radius: 4px;
                padding: 2px;
                font-size: 10pt;
                color: {var_theme.colors['text_primary']};
            }}
            QLineEdit:focus {{
                border: 2px solid {var_theme.colors['border_primary']};
            }}
        """)
        find_layout.addWidget(self.new_find_input)
        add_box_layout.addLayout(find_layout)
        replace_layout = QHBoxLayout()
        replace_label = QLabel("Replace:")
        replace_label.setMinimumWidth(55)
        replace_layout.addWidget(replace_label)
        self.new_replace_input = QLineEdit()
        self.new_replace_input.setPlaceholderText("Replacement text...")
        self.new_replace_input.setMinimumHeight(24)
        self.new_replace_input.setStyleSheet(f"""
            QLineEdit {{
                background-color: {var_theme.colors['input_bg']};
                border: 2px solid {var_theme.colors['border_primary']};
                border-radius: 4px;
                padding: 2px;
                font-size: 10pt;
                color: {var_theme.colors['text_primary']};
            }}
            QLineEdit:focus {{
                border: 2px solid {var_theme.colors['border_primary']};
            }}
        """)
        replace_layout.addWidget(self.new_replace_input)
        add_box_layout.addLayout(replace_layout)
        or_label = QLabel("— or use special character —")
        or_label.setStyleSheet("font-size: 8pt; font-style: italic; color: #aaaaaa; margin: 4px 0px;")
        or_label.setAlignment(Qt.AlignCenter)
        add_box_layout.addWidget(or_label)
        special_layout = QHBoxLayout()
        special_label = QLabel("Special:")
        special_label.setMinimumWidth(55)
        special_layout.addWidget(special_label)
        self.new_special_combo = QComboBox()
        self.new_special_combo.addItem("(None)", None)
        self.new_special_combo.addItem("Line break", "\n")
        self.new_special_combo.addItem("Space", " ")
        self.new_special_combo.addItem("Tab", "\t")
        self.new_special_combo.setMinimumHeight(24)
        self.new_special_combo.setStyleSheet(f"""
            QComboBox {{
                background-color: {var_theme.colors['input_bg']};
                border: 2px solid {var_theme.colors['border_primary']};
                border-radius: 4px;
                padding: 2px;
                font-size: 10pt;
                color: {var_theme.colors['text_primary']};
            }}
            QComboBox:focus {{
                border: 2px solid {var_theme.colors['border_primary']};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 25px;
            }}
            QComboBox::down-arrow {{
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 6px solid {var_theme.colors['text_secondary']};
            }}
        """)
        self.new_special_combo.currentIndexChanged.connect(
            lambda: self.on_special_replacement_changed(self.new_replace_input, self.new_special_combo)
        )
        special_layout.addWidget(self.new_special_combo)
        special_layout.addStretch()
        add_box_layout.addLayout(special_layout)
        add_box_layout.addSpacing(6)
        add_more_btn = QPushButton("+ Add Replacement")
        add_more_btn.setStyleSheet(get_button_style('success'))
        add_more_btn.setMinimumHeight(28)
        add_more_btn.setMaximumHeight(28)
        add_more_btn.clicked.connect(self.add_replacement_from_inputs)
        add_box_layout.addWidget(add_more_btn)
        add_box.setLayout(add_box_layout)
        add_box.setMinimumWidth(320)
        add_box.setMaximumWidth(320)
        add_box.setMinimumHeight(200)
        two_box_layout.addWidget(add_box)
        list_box = QGroupBox("Current Replacements")
        list_box.setStyleSheet(f"""
            QGroupBox {{
                background-color: {var_theme.colors['table_bg']};
                border: 2px solid {var_theme.colors['button_primary']};
                border-radius: 8px;
                font-weight: 600;
                font-size: 10pt;
                color: {var_theme.colors['button_primary']};
                padding-top: 18px;
                margin-top: 12px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top center;
                background-color: {var_theme.colors['table_bg']};
                padding: 4px 12px;
                border: 1px solid {var_theme.colors['button_primary']};
                border-radius: 4px;
            }}
        """)
        list_box_layout = QVBoxLayout()
        list_box_layout.setContentsMargins(8, 10, 8, 8)
        list_box_layout.setSpacing(6)
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(155)
        scroll_area.setMaximumHeight(155)
        scroll_area.setStyleSheet(f"""
            QScrollArea {{
                background-color: {var_theme.colors['window_bg']};
                border: 1px solid {var_theme.colors['border_dark']};
                border-radius: 4px;
            }}
        """)
        self.replacement_container = QWidget()
        self.replacement_container.setStyleSheet(f"background-color: {var_theme.colors['window_bg']};")
        self.replacement_layout = QVBoxLayout(self.replacement_container)
        self.replacement_layout.setSpacing(4)
        self.replacement_layout.setContentsMargins(6, 6, 6, 6)
        self.replacement_pairs = []
        self.replacement_layout.addStretch()
        scroll_area.setWidget(self.replacement_container)
        list_box_layout.addWidget(scroll_area)
        list_box.setLayout(list_box_layout)
        list_box.setMinimumWidth(680)
        list_box.setMaximumWidth(680)
        list_box.setMinimumHeight(200)
        two_box_layout.addWidget(list_box)
        replacement_section_layout.addLayout(two_box_layout)
        replacement_section.setLayout(replacement_section_layout)
        main_layout.addWidget(replacement_section)
        preview_group = QGroupBox("Formatting Preview")
        preview_layout = QVBoxLayout()
        preview_layout.setContentsMargins(6, 6, 6, 6)
        self.format_preview = QTextEdit()
        self.format_preview.setMinimumHeight(80)
        self.format_preview.setMaximumHeight(200)
        self.format_preview.setReadOnly(True)
        self.format_preview.setStyleSheet(
            """QTextEdit {
                background-color: #1e1e1e;
                color: #ffffff;
                border: 1px solid #555;
                font-family: 'Consolas', monospace;
                font-size: 9pt;
                padding: 0px;
            }"""
        )
        preview_layout.addWidget(self.format_preview)
        preview_group.setLayout(preview_layout)
        main_layout.addWidget(preview_group)
        nav_layout = QHBoxLayout()
        back_btn_template = QPushButton("← Back")
        back_btn_template.setStyleSheet(get_button_style('default'))
        back_btn_template.setMinimumHeight(32)
        back_btn_template.clicked.connect(lambda: self.tabs.setCurrentIndex(2))
        self.next_btn_template = QPushButton("Next: Send Emails →")
        self.next_btn_template.setStyleSheet(get_button_style('primary'))
        self.next_btn_template.setMinimumHeight(32)
        self.next_btn_template.setMinimumWidth(150)
        self.next_btn_template.clicked.connect(lambda: self.tabs.setCurrentIndex(4))
        nav_layout.addWidget(back_btn_template)
        nav_layout.addStretch()
        nav_layout.addWidget(self.next_btn_template)
        main_layout.addLayout(nav_layout)
        return widget
    def add_replacement_from_inputs(self):
        """Add a new replacement from the input fields"""
        find_text = self.new_find_input.text().strip()
        if not find_text:
            QMessageBox.warning(self, "Input Required", "Please enter text to find.")
            return
        replace_text = self.new_replace_input.text()
        special_value = self.new_special_combo.currentData()
        self.add_replacement_pair(find_text, replace_text, special_value)
        self.new_find_input.clear()
        self.new_replace_input.clear()
        self.new_special_combo.setCurrentIndex(0)
        self.new_replace_input.setEnabled(True)
        self.auto_save_and_update_preview()
    def add_replacement_pair(self, find_text="", replace_text="", special_value=None):
        """Add a new find/replace pair widget to the list"""
        pair_widget = QWidget()
        pair_layout = QHBoxLayout(pair_widget)
        pair_layout.setContentsMargins(8, 6, 8, 6)
        pair_layout.setSpacing(10)
        find_section = QLabel("Find:")
        find_section.setStyleSheet("color: #aaaaaa; font-size: 9pt;")
        find_section.setMinimumWidth(40)
        pair_layout.addWidget(find_section)
        find_label = QLabel(f'"{find_text}"')
        find_label.setStyleSheet("font-weight: bold; color: #4EC9B0; font-size: 9pt;")
        find_label.setMinimumWidth(120)
        find_label.setWordWrap(False)
        pair_layout.addWidget(find_label)
        arrow_label = QLabel("→")
        arrow_label.setStyleSheet("color: #ffffff; font-size: 10pt;")
        pair_layout.addWidget(arrow_label)
        replace_section = QLabel("Replace:")
        replace_section.setStyleSheet("color: #aaaaaa; font-size: 9pt;")
        replace_section.setMinimumWidth(55)
        pair_layout.addWidget(replace_section)
        if special_value is not None:
            special_name = {"\n": "Line break", " ": "Space", "\t": "Tab"}.get(special_value, "Special")
            replace_label = QLabel(f"[{special_name}]")
            replace_label.setStyleSheet("font-style: italic; color: #FFD700; font-size: 9pt;")
        else:
            replace_label = QLabel(f'"{replace_text}"')
            replace_label.setStyleSheet("color: #9CDCFE; font-size: 9pt;")
        replace_label.setMinimumWidth(120)
        replace_label.setWordWrap(False)
        pair_layout.addWidget(replace_label)
        pair_layout.addStretch()
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet(get_button_style('danger'))
        delete_btn.setMinimumWidth(75)
        delete_btn.setMinimumHeight(26)
        delete_btn.clicked.connect(lambda: self.remove_replacement_pair(pair_widget))
        pair_layout.addWidget(delete_btn)
        pair_data = {
            'widget': pair_widget,
            'find_text': find_text,
            'replace_text': replace_text,
            'special_value': special_value
        }
        self.replacement_pairs.append(pair_data)
        insert_position = self.replacement_layout.count() - 1
        self.replacement_layout.insertWidget(insert_position, pair_widget)
        return pair_data
    def remove_replacement_pair(self, pair_widget):
        """Reauto_save_and_updateeplace pair widget"""
        for i, pair_data in enumerate(self.replacement_pairs):
            if pair_data['widget'] == pair_widget:
                self.replacement_pairs.pop(i)
                break
        self.replacement_layout.removeWidget(pair_widget)
        pair_widget.deleteLater()
        self.update_format_preview()
    def on_special_replacement_changed(self, replace_input, special_combo):
        """Handle special replacement dropdown changes"""
        special_value = special_combo.currentData()
        if special_value is not None:
            replace_input.setEnabled(False)
            replace_input.setPlaceholderText(f"Using: {special_combo.currentText()}")
        else:
            replace_input.setEnabled(True)
            replace_input.setPlaceholderText("Replacement text...")
    def on_bullet_checkbox_changed(self, state):
        """Enable/disable bullet formatting options"""
        is_enabled = (state == 2)  
        self.bullet_combo.setEnabled(is_enabled)
        self.auto_save_and_update_preview()
    def load_formatting_settings(self):
        """Load saved formatting settings for selected column"""
        column = self.format_column_combo.currentText()
        if column in self.template_formatting:
            settings = self.template_formatting[column]
            bullet_enabled = settings.get('bullet_enabled', False)
            if hasattr(self, 'enable_bullet_checkbox'):
                self.enable_bullet_checkbox.setChecked(bullet_enabled)
            bullet = settings.get('bullet', '-')
            for i in range(self.bullet_combo.count()):
                if self.bullet_combo.itemData(i) == bullet:
                    self.bullet_combo.setCurrentIndex(i)
                    break
            replacements = settings.get('replacements', [])
            while self.replacement_pairs:
                pair = self.replacement_pairs[0]
                self.replacement_layout.removeWidget(pair['widget'])
                pair['widget'].deleteLater()
                self.replacement_pairs.pop(0)
            if replacements:
                for find_text, replace_text, special_type in replacements:
                    self.add_replacement_pair(find_text, replace_text, special_type)
        else:
            if hasattr(self, 'enable_bullet_checkbox'):
                self.enable_bullet_checkbox.setChecked(False)
            self.bullet_combo.setCurrentIndex(0)
            while self.replacement_pairs:
                pair = self.replacement_pairs[0]
                self.replacement_layout.removeWidget(pair['widget'])
                pair['widget'].deleteLater()
                self.replacement_pairs.pop(0)
        self.update_format_preview()
    def save_formatting_rules(self, show_message=True):
        """Save formatting rules for selected column"""
        column = self.format_column_combo.currentText()
        if not column:
            return
        replacements = []
        for pair in self.replacement_pairs:
            find_text = pair['find_text']
            replace_text = pair['replace_text']
            special_value = pair['special_value']
            if find_text:  
                replacements.append((find_text, replace_text, special_value))
        self.template_formatting[column] = {
            'bullet_enabled': self.enable_bullet_checkbox.isChecked() if hasattr(self, 'enable_bullet_checkbox') else False,
            'bullet': self.bullet_combo.currentData(),
            'replacements': replacements
        }
        if show_message:
            QMessageBox.information(self, "Saved", f"Formatting rules saved for column '{column}'")
    def auto_save_and_update_preview(self):
        """Automatically save formatting rules and update preview without showing message"""
        self.save_formatting_rules(show_message=False)
        self.update_format_preview()
    def update_format_preview(self):
        """Update the formatting preview with current settings"""
        try:
            column = self.format_column_combo.currentText()
            if not column or not hasattr(self, 'imported_data') or not self.imported_data:
                self.format_preview.setText("No data to preview")
                return
            col_index = None
            for i, header in enumerate(self.headers):
                if header == column:
                    col_index = i
                    break
            if col_index is None or not self.imported_data:
                self.format_preview.setText("No data found for this column")
                return
            preview_row_index = 0
            if self.selected_rows:
                preview_row_index = min(self.selected_rows)
                preview_source = f"Selected Row {preview_row_index + 1}"
            else:
                preview_source = "First Row (no selection)"
            if preview_row_index >= len(self.imported_data):
                preview_row_index = 0
            sample_data = self.imported_data[preview_row_index][col_index] if col_index < len(self.imported_data[preview_row_index]) else ""
            formatted = self.format_column_data_new(str(sample_data), column)
            self.format_preview.setText(f"Preview from: {preview_source}\n\nOriginal:\n{sample_data}\n\n{'='*40}\n\nFormatted:\n{formatted}")
        except Exception as e:
            self.format_preview.setText(f"Preview error: {str(e)}")
            logger.error(f"Error updating format preview: {e}")
    def format_column_data_new(self, data, column):
        """New formatting function with bullet points and string replacements"""
        if not data:
            return data
        formatted = str(data)
        if column not in self.template_formatting:
            return formatted
        settings = self.template_formatting[column]
        replacements = settings.get('replacements', [])
        for replacement_tuple in replacements:
            if len(replacement_tuple) == 3:
                find_text, replace_text, special_type = replacement_tuple
            else:
                find_text, replace_text = replacement_tuple
                special_type = None
            if find_text:
                actual_replacement = special_type if special_type is not None else replace_text
                formatted = formatted.replace(find_text, actual_replacement)
        bullet_enabled = settings.get('bullet_enabled', False)
        if bullet_enabled:
            bullet = settings.get('bullet', '-')
            lines = formatted.split('\n')
            formatted_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    formatted_lines.append(f"\t{bullet} {line}")
            formatted = '\n'.join(formatted_lines)
        return formatted
    def format_column_data(self, data, column):
        """Legacy compatibility - redirects to new formatting function"""
        return self.format_column_data_new(data, column)
    def process_template_placeholders(self, template, row_data):
        if not self.headers or not row_data:
            return template
        processed_template = template
        for i, header in enumerate(self.headers):
            placeholder = f"{{{header.upper()}}}"
            if placeholder in processed_template and i < len(row_data):
                data = str(row_data[i]) if row_data[i] is not None else ""
                if header in self.template_formatting:
                    data = self.format_column_data(data, header)
                processed_template = processed_template.replace(placeholder, data)
        return processed_template
    def create_send_tab(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(12)
        layout.setContentsMargins(15, 15, 15, 15)
        account_group = QGroupBox("Email Account")
        account_layout = QHBoxLayout()
        account_layout.setContentsMargins(12, 12, 12, 12)
        account_layout.setSpacing(8)
        self.account_combo = QComboBox()
        self.account_combo.setMinimumWidth(350)
        self.account_combo.setStyleSheet(var_theme.get_input_style())
        self.account_combo.currentIndexChanged.connect(self.on_account_changed)
        self.account_combo.currentIndexChanged.connect(self.update_send_summary)
        refresh_btn = QPushButton("Refresh")
        refresh_btn.setStyleSheet(get_button_style('default'))
        refresh_btn.clicked.connect(self.load_email_accounts)
        account_layout.addWidget(QLabel("Account:"))
        account_layout.addWidget(self.account_combo)
        account_layout.addWidget(refresh_btn)
        account_layout.addStretch()
        account_group.setLayout(account_layout)
        layout.addWidget(account_group)
        summary_group = QGroupBox("Send Summary")
        summary_layout = QVBoxLayout()
        summary_layout.setContentsMargins(12, 12, 12, 12)
        summary_layout.setSpacing(10)
        self.summary_label = QLabel("Configure settings above to see send summary")
        self.summary_label.setWordWrap(True)
        self.summary_label.setMinimumHeight(35)
        self.summary_label.setStyleSheet("padding: 10px; font-size: 10pt;")
        summary_layout.addWidget(self.summary_label)
        send_layout = QHBoxLayout()
        send_layout.addStretch()
        self.send_btn = QPushButton("Send Emails")
        self.send_btn.setStyleSheet(get_button_style('success'))
        self.send_btn.setMinimumHeight(45)
        self.send_btn.setMinimumWidth(160)
        self.send_btn.clicked.connect(self.send_emails)
        self.send_btn.setEnabled(False)
        send_layout.addWidget(self.send_btn)
        send_layout.addStretch()
        summary_layout.addLayout(send_layout)
        summary_group.setLayout(summary_layout)
        layout.addWidget(summary_group)
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        log_group = QGroupBox("Sending Log")
        log_layout = QVBoxLayout()
        self.log_display = QTextEdit()
        self.log_display.setMinimumHeight(180)
        self.log_display.setMaximumHeight(220)
        self.log_display.setReadOnly(True)
        self.log_display.setFont(QFont('Consolas', 9))
        self.log_display.setStyleSheet(var_theme.get_input_style())
        log_layout.addWidget(self.log_display)
        log_group.setLayout(log_layout)
        layout.addWidget(log_group)
        nav_layout = QHBoxLayout()
        back_btn_4 = QPushButton("← Back")
        back_btn_4.setStyleSheet(get_button_style('default'))
        back_btn_4.clicked.connect(lambda: self.tabs.setCurrentIndex(3))
        nav_layout.addWidget(back_btn_4)
        nav_layout.addStretch()
        layout.addLayout(nav_layout)
        return widget
    def apply_theme(self):
        self.setStyleSheet(var_theme.get_complete_style())
    def get_clickable_tab_style(self):
        """Return tab style with cursor pointer for clickable tabs"""
        def increase_saturation(hex_color, percent=10):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r, g, b = r/255.0, g/255.0, b/255.0
            max_c = max(r, g, b)
            min_c = min(r, g, b)
            l = (max_c + min_c) / 2.0
            if max_c == min_c:
                h = s = 0.0
            else:
                d = max_c - min_c
                s = d / (2.0 - max_c - min_c) if l > 0.5 else d / (max_c + min_c)
                if max_c == r:
                    h = (g - b) / d + (6.0 if g < b else 0.0)
                elif max_c == g:
                    h = (b - r) / d + 2.0
                else:
                    h = (r - g) / d + 4.0
                h /= 6.0
            s = min(1.0, s * (1.0 + percent/100.0))
            if s == 0:
                r = g = b = l
            else:
                def hue_to_rgb(p, q, t):
                    if t < 0: t += 1
                    if t > 1: t -= 1
                    if t < 1/6: return p + (q - p) * 6 * t
                    if t < 1/2: return q
                    if t < 2/3: return p + (q - p) * (2/3 - t) * 6
                    return p
                q = l * (1 + s) if l < 0.5 else l + s - l * s
                p = 2 * l - q
                r = hue_to_rgb(p, q, h + 1/3)
                g = hue_to_rgb(p, q, h)
                b = hue_to_rgb(p, q, h - 1/3)
            return f"#{int(r*255):02x}{int(g*255):02x}{int(b*255):02x}"
        selected_hover = increase_saturation(var_theme.colors['button_primary'], 10)
        unselected_hover = increase_saturation(var_theme.colors['secondary_bg'], 10)
        return f"""
            QTabWidget::pane {{
                background-color: {var_theme.colors['window_bg']};
                border: 1px solid {var_theme.colors['border_primary']};
                border-radius: 4px;
                margin: 2px;
            }}
            QTabBar::tab {{
                background-color: {var_theme.colors['secondary_bg']};
                border: 1px solid {var_theme.colors['border_primary']};
                border-bottom: none;
                border-radius: 4px 4px 0px 0px;
                color: {var_theme.colors['text_secondary']};
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 10pt;
                font-weight: 500;
                margin: 2px 1px 0px 1px;
                min-width: 180px;
                max-width: 200px;
                padding: 8px 6px;
            }}
            QTabBar::tab:selected {{
                background-color: {var_theme.colors['button_primary']};
                color: {var_theme.colors['text_inverse']};
                font-weight: 600;
            }}
            QTabBar::tab:hover:!selected:!disabled {{
                background-color: {unselected_hover};
                color: {var_theme.colors['text_primary']};
            }}
            QTabBar::tab:selected:hover:!disabled {{
                background-color: {selected_hover};
                color: {var_theme.colors['text_inverse']};
            }}
            QTabBar::tab:disabled {{
                background-color: {var_theme.colors['secondary_bg']};
                color: {var_theme.colors['text_muted']};
                border: 1px solid {var_theme.colors['border_light']};
            }}
        """
    def load_all_tabs(self):
        """Pre-load all tabs at startup with loading progress"""
        tab_labels = ["1. Import Data", "2. Compose Email", "3. Map Columns", 
                      "4. Template Formatting", "5. Send Emails"]
        tab_creators = [
            self.create_import_tab,
            self.create_compose_tab,
            self.create_mapping_tab,
            self.create_template_formatting_tab,
            self.create_send_tab
        ]
        for i, (label, creator) in enumerate(zip(tab_labels, tab_creators)):
            if self.loading_screen:
                progress = 40 + (i * 10)  
                self.loading_screen.update_progress(progress, f"Loading {label}...")
                QApplication.processEvents()
            widget = creator()
            self.tabs.addTab(widget, label)
            self.tabs_created.add(i)
            self.tab_widgets[i] = widget
        if self.loading_screen:
            self.loading_screen.update_progress(85, "Loading email accounts...")
            QApplication.processEvents()
        if not self.email_accounts_loaded:
            self.load_email_accounts()
            self.email_accounts_loaded = True
        self.tabs.setCurrentIndex(0)
    def on_tab_clicked(self, index):
        """Handle tab click navigation"""
        if self.tabs.isTabEnabled(index):
            self.tabs.setCurrentIndex(index)
    def on_tab_changed(self, index):
        """Handle tab switching - reload data when needed"""
        # Reload email accounts when switching to Send tab (index 4)
        if index == 4 and hasattr(self, 'account_combo'):
            self.load_email_accounts()
        
        # Update mapping table when switching to Mapping tab (index 2)
        if index == 2 and hasattr(self, 'mapping_table'):
            self.update_mapping_table()
        
        # Update format preview when switching to Template Formatting tab (index 3)
        if index == 3 and hasattr(self, 'format_preview'):
            self.update_format_preview()  
    def browse_file(self):
        file_filter = "All Files (*.*);;CSV Files (*.csv);;Excel Files (*.xlsx)"
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select File to Import", "", file_filter
        )
        if file_path:
            self.file_path_input.setText(file_path)
            self.import_btn.setEnabled(True)
    def import_file(self):
        file_path = self.file_path_input.text().strip()
        if not file_path:
            QMessageBox.warning(self, "Warning", "Please select a file to import.")
            return
        try:
            result = FileImporter.import_file(file_path)
            if result['success']:
                self.imported_data = result['data']
                self.headers = result['headers']
                if hasattr(self, 'format_column_combo'):
                    self.format_column_combo.clear()
                    self.format_column_combo.addItems(self.headers)
                self.populate_data_table()
                self.next_btn_1.setEnabled(True)
                self.statusBar().showMessage(f"Imported {len(self.imported_data)} rows")
                # Instantly update mapping table if placeholders exist
                if self.placeholders and hasattr(self, 'mapping_table'):
                    self.update_mapping_table()
                QMessageBox.information(self, "Success", result['message'])
            else:
                QMessageBox.critical(self, "Import Error", result['message'])
        except Exception as e:
            logger.error(f"Critical error during import: {e}")
            QMessageBox.critical(self, "Critical Error", f"An unexpected error occurred: {str(e)}")
    def populate_data_table(self):
        if not self.imported_data or not self.headers:
            return
        self.filtered_data = self.imported_data.copy()
        self.filter_column_combo.clear()
        self.filter_column_combo.addItem("-- Select Column --")
        self.filter_column_combo.addItems(self.headers)
        self.update_table_display()
        self.data_table.selectionModel().selectionChanged.connect(self.update_selection_info)
        self.update_selection_info()
    def update_table_display(self):
        try:
            self.data_table.itemChanged.disconnect(self.on_checkbox_changed)
        except:
            pass
        display_data = self.filtered_data  
        self.data_table.setRowCount(len(display_data))
        self.data_table.setColumnCount(len(self.headers) + 1)  
        table_headers = ["✓"] + self.headers
        self.data_table.setHorizontalHeaderLabels(table_headers)
        self.data_table.setColumnWidth(0, 40)
        for row, data in enumerate(display_data):
            checkbox_item = QTableWidgetItem()
            checkbox_item.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled)
            original_row_index = -1
            for idx, imported_row in enumerate(self.imported_data):
                if len(imported_row) == len(data) and all(str(a) == str(b) for a, b in zip(imported_row, data)):
                    original_row_index = idx
                    break
            if original_row_index in self.selected_rows:
                checkbox_item.setCheckState(Qt.Checked)
            else:
                checkbox_item.setCheckState(Qt.Unchecked)
            self.data_table.setItem(row, 0, checkbox_item)
            for col, value in enumerate(data):
                item = QTableWidgetItem(str(value) if value is not None else '')
                item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)  
                self.data_table.setItem(row, col + 1, item)  
        self.data_table.itemChanged.connect(self.on_checkbox_changed)
        self.data_table.resizeColumnsToContents()
    def update_selection_info(self):
        checked_rows = len(self.selected_rows)
        total_filtered = len(self.filtered_data)
        total_imported = len(self.imported_data)
        if total_imported == 0:
            info_text = "No data loaded"
        elif total_filtered == total_imported:
            info_text = f"Showing all {total_filtered} rows | Checked: {checked_rows} rows"
        else:
            info_text = f"Showing {total_filtered} filtered rows (Total: {total_imported}) | Checked: {checked_rows} rows"
        self.selection_info_label.setText(info_text)
        # Instantly update send summary when selection changes
        self.update_send_summary()
    def filter_table_data(self):
        if not self.imported_data:
            return
        search_text = self.search_input.text().strip().lower()
        if not search_text:
            self.filtered_data = self.imported_data.copy()
        else:
            self.filtered_data = []
            for row in self.imported_data:
                row_text = ' '.join(str(cell).lower() for cell in row if cell is not None)
                if search_text in row_text:
                    self.filtered_data.append(row)
        self.sort_table_data()
        self.update_table_display()
        self.update_selection_info()
    def sort_table_data(self):
        if not self.filtered_data or not hasattr(self, 'filter_column_combo'):
            return
        column_name = self.filter_column_combo.currentText()
        if column_name == "-- Select Column --" or column_name not in self.headers:
            return
        column_index = self.headers.index(column_name)
        sort_order = self.sort_order_combo.currentText()
        reverse = (sort_order == "Z-A")
        try:
            self.filtered_data.sort(
                key=lambda row: str(row[column_index] if column_index < len(row) and row[column_index] is not None else ''),
                reverse=reverse
            )
            self.update_table_display()
            self.update_selection_info()
        except Exception as e:
            logger.warning(f"Error sorting data: {e}")
    def clear_filters(self):
        self.search_input.clear()
        self.filter_column_combo.setCurrentIndex(0)
        self.sort_order_combo.setCurrentIndex(0)
        self.filtered_data = self.imported_data.copy()
        self.update_table_display()
        self.update_selection_info()
    def on_checkbox_changed(self, item):
        if item.column() == 0:  
            row = item.row()
            if row < len(self.filtered_data):
                data_row = self.filtered_data[row]
                original_index = -1
                for idx, imported_row in enumerate(self.imported_data):
                    if len(imported_row) == len(data_row) and all(str(a) == str(b) for a, b in zip(imported_row, data_row)):
                        original_index = idx
                        break
                if original_index >= 0:
                    if item.checkState() == Qt.Checked:
                        self.selected_rows.add(original_index)
                    else:
                        self.selected_rows.discard(original_index)
            self.update_selection_info()
            if hasattr(self, 'format_preview'):
                self.update_format_preview()
    def select_all_rows(self):
        try:
            self.data_table.itemChanged.disconnect(self.on_checkbox_changed)
        except:
            pass
        for row in range(self.data_table.rowCount()):
            checkbox_item = self.data_table.item(row, 0)
            if checkbox_item:
                checkbox_item.setCheckState(Qt.Checked)
                if row < len(self.filtered_data):
                    data_row = self.filtered_data[row]
                    for idx, imported_row in enumerate(self.imported_data):
                        if len(imported_row) == len(data_row) and all(str(a) == str(b) for a, b in zip(imported_row, data_row)):
                            self.selected_rows.add(idx)
                            break
        self.data_table.itemChanged.connect(self.on_checkbox_changed)
        self.update_selection_info()
        if hasattr(self, 'format_preview'):
            self.update_format_preview()
    def deselect_all_rows(self):
        try:
            self.data_table.itemChanged.disconnect(self.on_checkbox_changed)
        except:
            pass
        for row in range(self.data_table.rowCount()):
            checkbox_item = self.data_table.item(row, 0)
            if checkbox_item:
                checkbox_item.setCheckState(Qt.Unchecked)
        self.selected_rows.clear()
        self.data_table.itemChanged.connect(self.on_checkbox_changed)
        self.update_selection_info()
        if hasattr(self, 'format_preview'):
            self.update_format_preview()
    def load_default_template(self):
        """Method kept for backwards compatibility but does nothing.
        Users must load template from file using the Load button."""
        pass
    def load_template(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Load Template", "", "Text Files (*.txt);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    template = f.read()
                self.template_editor.setPlainText(template)
                self.detect_placeholders()
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error loading template: {e}")
    def save_template(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Template", "email_template.txt", "Text Files (*.txt);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.template_editor.toPlainText())
                QMessageBox.information(self, "Success", "Template saved successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error saving template: {e}")
    def detect_placeholders(self):
        """Detect placeholders in the template - handles lazy-loaded UI elements"""
        try:
            if not hasattr(self, 'template_editor'):
                return
            template_text = self.template_editor.toPlainText()
            self.placeholders = PlaceholderExtractor.extract_placeholders(template_text)
            if hasattr(self, 'placeholders_label'):
                if self.placeholders:
                    placeholder_text = f"Detected placeholders: {', '.join(self.placeholders)}"
                    color = var_theme.colors['success']
                    if hasattr(self, 'next_btn_2'):
                        self.next_btn_2.setEnabled(True)
                else:
                    placeholder_text = "No placeholders detected. Use {NAME}, <EMAIL>, [[DEVICES]], etc."
                    color = var_theme.colors['warning']
                self.placeholders_label.setText(placeholder_text)
                self.placeholders_label.setStyleSheet(
                    f"color: {color}; padding: 8px; "
                    f"background-color: {var_theme.colors['secondary_bg']}; "
                    f"border-radius: 4px; font-size: 9pt;"
                )
            if self.headers and hasattr(self, 'mapping_table'):
                self.update_mapping_table()
        except Exception as e:
            logger.error(f"Error detecting placeholders: {e}")
            import traceback
            logger.error(traceback.format_exc())
            if hasattr(self, 'placeholders_label'):
                self.placeholders_label.setText("Error detecting placeholders")
                self.placeholders_label.setStyleSheet(f"color: {var_theme.colors['error']}; padding: 8px;")
    def update_mapping_table(self):
        """Update the mapping table - handles lazy-loaded UI elements"""
        if not self.placeholders:
            return
        if not hasattr(self, 'mapping_table'):
            return
        self.mapping_table.setRowCount(len(self.placeholders))
        suggestions = PlaceholderExtractor.suggest_mappings(self.placeholders, self.headers)
        for row, placeholder in enumerate(self.placeholders):
            placeholder_item = QTableWidgetItem(placeholder)
            placeholder_item.setFlags(Qt.ItemIsEnabled)
            self.mapping_table.setItem(row, 0, placeholder_item)
            column_combo = QComboBox()
            column_combo.addItem("-- Select Column --")
            column_combo.addItems(self.headers)
            column_combo.setMinimumWidth(260)  
            column_combo.setMaximumWidth(260)
            column_combo.setMinimumHeight(25)   
            column_combo.setStyleSheet(var_theme.get_input_style() + 
                f"""QComboBox {{
                    padding: 4px 8px;
                    font-size: 10pt;
                }}
                QComboBox::drop-down {{
                    width: 25px;
                }}
                QComboBox::item {{
                    padding: 4px 8px;
                    min-height: 20px;
                }}""")
            column_combo.currentTextChanged.connect(self.update_send_summary)
            if placeholder in suggestions:
                suggested_column = suggestions[placeholder]
                if suggested_column in self.headers:
                    index = self.headers.index(suggested_column) + 1
                    column_combo.setCurrentIndex(index)
            self.mapping_table.setCellWidget(row, 1, column_combo)
            sample_data = ""
            if placeholder in suggestions and suggestions[placeholder] in self.headers:
                col_index = self.headers.index(suggestions[placeholder])
                if self.imported_data and col_index < len(self.imported_data[0]):
                    sample_data = str(self.imported_data[0][col_index])[:50]
            sample_item = QTableWidgetItem(sample_data)
            sample_item.setFlags(Qt.ItemIsEnabled)
            self.mapping_table.setItem(row, 2, sample_item)
        self.mapping_table.setColumnWidth(0, 200)  
        self.mapping_table.setColumnWidth(1, 280)  
        self.mapping_table.setColumnWidth(2, 650)  
        self.mapping_table.resizeRowsToContents()
    def on_account_changed(self, index):
        """Log when user changes the selected account"""
        if index >= 0 and index < len(self.email_accounts_list):
            selected_email = self.email_accounts_list[index]['email']
            logger.info(f"\n>>> USER SELECTED: Index[{index}] = {selected_email} <<<\n")
            if hasattr(self, 'log_display'):
                self.log_display.append(f"Selected sender: {selected_email}")
    def load_email_accounts(self):
        """Load email accounts from Outlook - handles lazy-loaded UI elements"""
        if not hasattr(self, 'account_combo'):
            logger.warning("Account combo not loaded yet, skipping email account loading")
            return
        self.account_combo.clear()
        self.email_accounts_list.clear()
        self.account_combo.addItem("No Outlook Integration Available")
        self.email_accounts = []
        self.email_accounts_list = []
        try:
            detected_accounts = EmailSender.get_email_accounts()
            self.account_combo.clear()
            self.email_accounts_list = detected_accounts.copy()  
            self.email_accounts = detected_accounts  
            if self.email_accounts_list:
                logger.info(f"\n{'='*60}")
                logger.info(f"DROPDOWN MENU ACCOUNT MAPPING:")
                for index, account in enumerate(self.email_accounts_list):
                    account_display = f"{account['email']} (Account {index + 1})"
                    self.account_combo.addItem(account_display)
                    self.account_combo.setItemData(index, account['email'])
                    logger.info(f"  Dropdown Index[{index}] → {account['email']}")
                logger.info(f"{'='*60}\n")
            else:
                self.account_combo.addItem("No email accounts found")
                self.email_accounts_list = []
                self.email_accounts = []
        except Exception as e:
            logger.error(f"Error loading email accounts: {e}")
            import traceback
            logger.error(traceback.format_exc())
            self.account_combo.addItem("Error loading accounts")
            self.email_accounts_list = []
            self.email_accounts = []
        self.update_send_summary()
    def update_send_summary(self, *args):
        if not self.imported_data:
            if hasattr(self, 'summary_label'):
                self.summary_label.setText("No data imported")
            if hasattr(self, 'send_btn'):
                self.send_btn.setEnabled(False)
            return
        recipient_count = len(self.selected_rows)  
        has_subject = bool(self.subject_input.text().strip()) if hasattr(self, 'subject_input') else False
        has_template = bool(self.template_editor.toPlainText().strip()) if hasattr(self, 'template_editor') else False
        has_account = (bool(self.email_accounts_list) and 
                      hasattr(self, 'account_combo') and 
                      self.account_combo.currentIndex() >= 0 and
                      self.account_combo.currentIndex() < len(self.email_accounts_list)) if hasattr(self, 'account_combo') else False
        if hasattr(self, 'next_btn_2'):
            self.next_btn_2.setEnabled(has_subject and has_template)
        if hasattr(self, 'next_btn_3'):
            self.next_btn_3.setEnabled(has_subject and has_template)
        if hasattr(self, 'next_btn_template'):
            self.next_btn_template.setEnabled(True)
        if hasattr(self, 'summary_label'):
            if has_subject and has_template and has_account and recipient_count > 0:
                account_index = self.account_combo.currentIndex()
                account_email = self.email_accounts_list[account_index]['email'] if self.email_accounts_list and account_index < len(self.email_accounts_list) else "Unknown"
                self.summary_label.setText(
                    f"Ready to send emails!\n\n"
                    f"• Recipients: {recipient_count} selected contacts\n"
                    f"• Subject: ✓\n"
                    f"• Template: ✓\n" 
                    f"• Sending from: {account_email}"
                )
            else:
                summary_parts = []
                if recipient_count == 0:
                    summary_parts.append("Recipients: None selected")
                else:
                    summary_parts.append(f"Recipients: {recipient_count} selected")
                summary_parts.extend([
                    f"Subject: {'✓' if has_subject else '✗'}",
                    f"Template: {'✓' if has_template else '✗'}",
                    f"Account: {'✓' if has_account else '✗'}"
                ])
                self.summary_label.setText(" | ".join(summary_parts))
        send_ready = has_subject and has_template and has_account and recipient_count > 0
        if hasattr(self, 'send_btn'):
            self.send_btn.setEnabled(send_ready)
    def send_emails(self):
        try:
            subject = self.subject_input.text().strip()
            template = self.template_editor.toPlainText().strip()
            if not subject or not template:
                QMessageBox.warning(self, "Missing Information", 
                                  "Please provide both email subject and template.")
                return
            if not self.email_accounts_list or self.account_combo.currentIndex() < 0 or self.account_combo.currentIndex() >= len(self.email_accounts_list):
                QMessageBox.warning(self, "No Account", 
                                  "Please select a valid email account.")
                return
            if not self.selected_rows:
                QMessageBox.warning(self, "No Recipients Selected", 
                                  "Please select at least one recipient by checking the boxes in the data table.")
                return
            try:
                mappings = {}
                for row in range(self.mapping_table.rowCount()):
                    placeholder = self.mapping_table.item(row, 0).text()
                    combo = self.mapping_table.cellWidget(row, 1)
                    if combo and combo.currentIndex() > 0:
                        column_name = combo.currentText()
                        if column_name in self.headers:
                            col_index = self.headers.index(column_name)
                            mappings[placeholder] = col_index
            except Exception as e:
                logger.error(f"Error building mappings: {e}")
                QMessageBox.critical(self, "Mapping Error", f"Error processing column mappings: {str(e)}")
                return
            try:
                recipients = []
                for row_index in sorted(self.selected_rows):
                    if row_index < len(self.imported_data):
                        row_data = self.imported_data[row_index]
                        recipient = {}
                        for header_index, header in enumerate(self.headers):
                            if header_index < len(row_data):
                                recipient[header] = row_data[header_index]
                        recipients.append(recipient)
                if not recipients:
                    QMessageBox.warning(self, "No Valid Recipients", 
                                      "No valid recipients found in selected rows.")
                    return
            except Exception as e:
                logger.error(f"Error building recipients: {e}")
                QMessageBox.critical(self, "Recipients Error", f"Error processing recipients: {str(e)}")
                return
            reply = QMessageBox.question(
                self, "Confirm Sending",
                f"Send emails to {len(recipients)} recipients?" + 
                (f"\nAttachments: {len(self.attachments)} files" if self.attachments else ""),
                QMessageBox.Yes | QMessageBox.No
            )
            if reply != QMessageBox.Yes:
                return
            account_index = self.account_combo.currentIndex()
            if account_index < 0 or account_index >= len(self.email_accounts_list):
                QMessageBox.warning(self, "Account Error", "Please select a valid email account.")
                return
            selected_account = self.email_accounts_list[account_index]
            sender_email = selected_account['email']
            logger.info(f"Selected dropdown index[{account_index}] = Account {account_index + 1}")
            logger.info(f"Sender email address: {sender_email}")
            self.log_display.append(f"Using sender account: {sender_email}")
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            self.send_btn.setEnabled(False)
            QApplication.processEvents()
            try:
                processed_recipients = []
                for recipient in recipients:
                    email_body = template
                    email_subject = subject
                    for header in self.headers:
                        placeholders = [
                            f"{{{header.upper()}}}",
                            f"{{{header}}}",
                            f"<{header.upper()}>",
                            f"<{header}>"
                        ]
                        if header in recipient:
                            data = str(recipient[header]) if recipient[header] is not None else ""
                            if header in self.template_formatting:
                                data = self.format_column_data(data, header)
                            for placeholder in placeholders:
                                email_body = email_body.replace(placeholder, data)
                                email_subject = email_subject.replace(placeholder, data)
                    processed_recipient = recipient.copy()
                    processed_recipient['_processed_template'] = email_body
                    processed_recipient['_processed_subject'] = email_subject
                    processed_recipients.append(processed_recipient)
                logger.info(f"Sending {len(processed_recipients)} emails from: {sender_email}")
                result = EmailSender.send_emails(
                    processed_recipients, subject, template,  
                    selected_account, self.attachments
                )
                self.progress_bar.setVisible(False)
                self.send_btn.setEnabled(True)
                if result['success']:
                    QMessageBox.information(
                        self, "Success",
                        f"Successfully sent {result['sent']} emails!"
                    )
                    self.log_display.append(f"✓ Successfully sent {result['sent']} emails")
                else:
                    failed_details = ""
                    if 'failed_details' in result and result['failed_details']:
                        failed_details = "\\n".join(result['failed_details'][:5])
                        if len(result['failed_details']) > 5:
                            failed_details += f"\\n... and {len(result['failed_details']) - 5} more"
                    QMessageBox.warning(
                        self, "Partial Success",
                        f"Sent: {result['sent']} emails\\n"
                        f"Failed: {result['failed']} emails\\n\\n"
                        f"Failed recipients:\\n{failed_details}"
                    )
                    self.log_display.append(f"⚠ Sent {result['sent']}, Failed {result['failed']} emails")
            except Exception as e:
                self.progress_bar.setVisible(False)
                self.send_btn.setEnabled(True)
                QMessageBox.critical(
                    self, "Error",
                    f"Error sending emails: {str(e)}"
                )
                self.log_display.append(f"✗ Error: {str(e)}")
        except Exception as e:
            logger.error(f"Critical error in send_emails: {e}")
            QMessageBox.critical(self, "Critical Error", f"A critical error occurred: {str(e)}")
    def add_attachment(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, "Select Attachments", "", 
            "All Files (*);;Documents (*.pdf *.doc *.docx);;Images (*.jpg *.jpeg *.png *.gif);;Excel Files (*.xlsx *.xls)"
        )
        if file_paths:
            for file_path in file_paths:
                if file_path not in self.attachments:
                    self.attachments.append(file_path)
            self.update_attachments_display()
            for file_path in file_paths:
                if file_path not in self.attachments:
                    self.attachments.append(file_path)
            self.update_attachments_display()
    def remove_attachment(self):
        current_row = self.attachments_table.currentRow()
        if current_row >= 0 and current_row < len(self.attachments):
            self.attachments.pop(current_row)
            self.update_attachments_display()
            self.attachments.pop(current_row)
            self.update_attachments_display()
    def clear_attachments(self):
        self.attachments.clear()
        self.update_attachments_display()
    def update_attachments_display(self):
        self.attachments_table.setRowCount(len(self.attachments))
        for row, attachment in enumerate(self.attachments):
            filename = os.path.basename(attachment)
            filename_item = QTableWidgetItem(filename)
            filename_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
            self.attachments_table.setItem(row, 0, filename_item)
            try:
                size_bytes = os.path.getsize(attachment)
                if size_bytes < 1024:
                    size_text = f"{size_bytes} B"
                elif size_bytes < 1024*1024:
                    size_text = f"{size_bytes/1024:.1f} KB"
                else:
                    size_text = f"{size_bytes/(1024*1024):.1f} MB"
            except OSError:
                size_text = "Unknown"
            size_item = QTableWidgetItem(size_text)
            size_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
            self.attachments_table.setItem(row, 1, size_item)
        if self.attachments:
            total_size = 0
            for attachment in self.attachments:
                try:
                    total_size += os.path.getsize(attachment)
                except OSError:
                    pass
            size_mb = total_size / (1024 * 1024)
            self.attachment_info_label.setText(
                f"{len(self.attachments)} file(s) ({size_mb:.1f}MB)"
            )
            self.attachment_info_label.setStyleSheet(f"color: {var_theme.colors['success']}; font-size: 8pt; padding: 2px;")
        else:
            self.attachment_info_label.setText("No files")
            self.attachment_info_label.setStyleSheet(f"color: {var_theme.colors['text_muted']}; font-size: 8pt; padding: 2px;")
    def closeEvent(self, event):
        """Handle window close event"""
        logger.info("Application closing...")
        event.accept()
    def apply_dark_titlebar(self):
        """Apply dark theme to Windows title bar using DWM API"""
        try:
            import ctypes
            from ctypes import wintypes
            hwnd = int(self.winId())
            DWMWA_USE_IMMERSIVE_DARK_MODE = 20
            value = ctypes.c_int(1)
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                hwnd,
                DWMWA_USE_IMMERSIVE_DARK_MODE,
                ctypes.byref(value),
                ctypes.sizeof(value)
            )
        except Exception as e:
            pass  
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setApplicationName("Universal Email Sender")
    from theme import apply_theme
    apply_theme(app)
    window = UniversalSender()
    window.show()
    sys.exit(app.exec_())